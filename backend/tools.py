"""
Tools para el Agente SGD·IA
Herramientas formales que el agente puede usar
"""

import os
import uuid
import re
import tempfile
from datetime import datetime
from typing import List, Dict, Any, Optional
from langchain.tools import tool
import pandas as pd
import numpy as np

# Importaciones del sistema existente
from backend.rag import (
    collection, get_documents_info, retrieve_chunks,
    get_conversation_context, save_to_memory
)


@tool
def buscar_en_documentos(query: str) -> str:
    """
    Busca información en los documentos indexados.
    
    Args:
        query: La pregunta o consulta del usuario
    
    Returns:
        Fragmentos relevantes encontrados en los documentos
    
    Ejemplos:
        - buscar_en_documentos("¿Qué dice el contrato sobre los plazos?")
        - buscar_en_documentos("¿Cuáles son los integrantes del proyecto?")
    """
    print(f"[TOOL] Buscando en documentos: {query[:50]}...")
    
    try:
        from backend.rag import retrieve_chunks, build_context
        
        docs, metas = retrieve_chunks(query, None, 15)
        
        if not docs:
            return "No se encontró información relevante en los documentos cargados."
        
        context, sources, _ = build_context(docs, metas, query, None)
        
        if len(context) > 3000:
            context = context[:3000] + "...\n\n[Contenido truncado por longitud]"
        
        fuentes = list(set([s.get("filename", "") for s in sources[:5]]))
        resultado = f"{context}\n\n**Fuentes:** {', '.join(fuentes)}"
        
        return resultado
        
    except Exception as e:
        print(f"[TOOL] Error en buscar_en_documentos: {e}")
        return f"Error al buscar en documentos: {str(e)}"


@tool
def resumir_documento(nombre_documento: str) -> str:
    """
    Genera un resumen automático de un documento específico.
    
    Args:
        nombre_documento: Nombre del documento a resumir (ej: "contrato.pdf")
    
    Returns:
        Resumen del contenido del documento
    
    Ejemplos:
        - resumir_documento("FORMATO_INFORME_AVANCE_EIE.pdf")
        - resumir_documento("Tipos_de_Lenguajes_SQL.docx")
    """
    print(f"[TOOL] Resumiendo documento: {nombre_documento}")
    
    try:
        from backend.rag import retrieve_chunks, build_context, _get_gemini_model
        
        docs, metas = retrieve_chunks(f"todo el contenido del documento {nombre_documento}", None, 20)
        
        if not docs:
            return f"No se encontró el documento '{nombre_documento}'. Verifica el nombre."
        
        chunks_doc = []
        for doc, meta in zip(docs, metas):
            if nombre_documento.lower() in meta.get("filename", "").lower():
                chunks_doc.append((doc, meta))
        
        if not chunks_doc:
            return f"No se encontró contenido para '{nombre_documento}'."
        
        texto_completo = "\n".join([d[:500] for d, _ in chunks_doc[:10]])
        
        model = _get_gemini_model()
        prompt = f"""
        Genera un resumen CONCISO del siguiente documento:
        
        DOCUMENTO: {nombre_documento}
        
        CONTENIDO:
        {texto_completo[:4000]}
        
        El resumen debe incluir:
        1. Tema principal
        2. Puntos clave (3-5)
        3. Si es aplicable, fechas, autores o datos importantes
        
        Resumen (máximo 10 líneas):
        """
        
        response = model.generate_content(prompt)
        resumen = response.text
        
        return f"**Resumen de '{nombre_documento}':**\n\n{resumen}"
        
    except Exception as e:
        print(f"[TOOL] Error en resumir_documento: {e}")
        return f"Error al resumir el documento: {str(e)}"


@tool
def extraer_datos_numericos(nombre_documento: str) -> str:
    """
    Extrae datos numéricos de un documento Excel o CSV.
    
    Args:
        nombre_documento: Nombre del archivo Excel o CSV
    
    Returns:
        Lista de datos numéricos encontrados y estadísticas
    
    Ejemplos:
        - extraer_datos_numericos("ventas_2024.xlsx")
        - extraer_datos_numericos("datos.csv")
    """
    print(f"[TOOL] Extrayendo datos numéricos de: {nombre_documento}")
    
    try:
        from backend.rag import collection
        import re
        
        all_meta = collection.get()["metadatas"]
        doc_id = None
        
        for meta in all_meta:
            if nombre_documento.lower() in meta.get("filename", "").lower():
                doc_id = meta.get("doc_id")
                break
        
        if not doc_id:
            return f"No se encontró el documento '{nombre_documento}'."
        
        chunks = collection.get(where={"doc_id": str(doc_id)})
        
        numeros = []
        for chunk in chunks.get("documents", []):
            encontrados = re.findall(r'-?\d+\.?\d*', chunk)
            for n in encontrados:
                try:
                    num = float(n)
                    if 0 < num < 10000000:
                        numeros.append(num)
                except:
                    pass
        
        if not numeros:
            return f"No se encontraron datos numéricos en '{nombre_documento}'."
        
        numeros_ordenados = sorted(numeros)
        total = len(numeros)
        suma = sum(numeros)
        promedio = suma / total if total > 0 else 0
        minimo = min(numeros)
        maximo = max(numeros)
        
        muestra = numeros[:20]
        
        resultado = f"""📊 **Datos extraídos de '{nombre_documento}':**

**Estadísticas:**
- Total de valores: {total}
- Suma total: {suma:,.2f}
- Promedio: {promedio:.2f}
- Mínimo: {minimo:.2f}
- Máximo: {maximo:.2f}

**Primeros valores:** {', '.join([str(round(v, 2)) for v in muestra[:10]])}

**Sugerencia:** Puedes pedirme que genere una gráfica con estos datos."""
        
        return resultado
        
    except Exception as e:
        print(f"[TOOL] Error en extraer_datos_numericos: {e}")
        return f"Error al extraer datos: {str(e)}"


@tool
def generar_word(contenido: str, titulo: str = "Documento SGD·IA") -> str:
    """
    Genera un archivo Word descargable con el contenido proporcionado.
    
    Args:
        contenido: El texto o información a incluir en el Word
        titulo: Título del documento (opcional)
    
    Returns:
        URL de descarga del archivo generado
    
    Ejemplos:
        - generar_word(resumen_contrato, "Resumen Contrato")
    """
    print(f"[TOOL] Generando Word: {titulo}")
    
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        filename = f"word_{uuid.uuid4().hex[:8]}.docx"
        temp_path = os.path.join(tempfile.gettempdir(), filename)
        
        doc = DocxDocument()
        
        title = doc.add_heading(titulo, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Documento: {titulo}")
        doc.add_page_break()
        
        doc.add_heading('Contenido', level=1)
        for parrafo in contenido.split('\n'):
            if parrafo.strip():
                doc.add_paragraph(parrafo)
        
        doc.save(temp_path)
        
        url = f"http://localhost:8000/api/download/{filename}"
        return f"✅ Word generado. **Descarga:** {url} (válido por 5 minutos)"
        
    except Exception as e:
        print(f"[TOOL] Error en generar_word: {e}")
        return f"Error al generar Word: {str(e)}"


@tool
def generar_pdf(contenido: str, titulo: str = "Documento SGD·IA") -> str:
    """
    Genera un archivo PDF descargable con el contenido proporcionado.
    
    Args:
        contenido: El texto o información a incluir en el PDF
        titulo: Título del documento (opcional)
    
    Returns:
        URL de descarga del archivo generado
    
    Ejemplos:
        - generar_pdf(informe_ventas, "Informe Ventas")
    """
    print(f"[TOOL] Generando PDF: {titulo}")
    
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        
        filename = f"pdf_{uuid.uuid4().hex[:8]}.pdf"
        temp_path = os.path.join(tempfile.gettempdir(), filename)
        
        doc = SimpleDocTemplate(temp_path, pagesize=letter)
        styles = getSampleStyleSheet()
        
        story = []
        story.append(Paragraph(titulo, styles['Heading1']))
        story.append(Spacer(1, 12))
        story.append(Paragraph(f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        story.append(Spacer(1, 12))
        story.append(Paragraph("Contenido", styles['Heading2']))
        
        for line in contenido.split('\n'):
            if line.strip():
                story.append(Paragraph(line.replace('\n', '<br/>'), styles['Normal']))
                story.append(Spacer(1, 6))
        
        doc.build(story)
        
        url = f"http://localhost:8000/api/download/{filename}"
        return f"✅ PDF generado. **Descarga:** {url} (válido por 5 minutos)"
        
    except Exception as e:
        print(f"[TOOL] Error en generar_pdf: {e}")
        return f"Error al generar PDF: {str(e)}"


@tool
def generar_grafica(datos: str, tipo: str = "auto", titulo: str = "Gráfica") -> str:
    """
    Genera una gráfica a partir de datos numéricos.
    
    Args:
        datos: Lista de números separados por comas (ej: "10,20,30,40")
        tipo: Tipo de gráfica ("lineas", "barras", "pastel", "auto")
        titulo: Título de la gráfica
    
    Returns:
        URL de descarga de la imagen
    
    Ejemplos:
        - generar_grafica("10,20,30,40", "barras", "Ventas 2024")
    """
    print(f"[TOOL] Generando gráfica: {titulo}")
    
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        
        datos_lista = [float(x.strip()) for x in datos.split(',')]
        
        if not datos_lista:
            return "No se proporcionaron datos válidos para la gráfica."
        
        filename = f"grafica_{uuid.uuid4().hex[:8]}.png"
        temp_path = os.path.join(tempfile.gettempdir(), filename)
        
        if tipo == "auto":
            if len(datos_lista) <= 5:
                tipo = "pastel"
            elif len(set(datos_lista)) < len(datos_lista) * 0.7:
                tipo = "barras"
            else:
                tipo = "lineas"
        
        plt.figure(figsize=(10, 6))
        
        if tipo == "lineas":
            plt.plot(datos_lista, marker='o', linewidth=2, markersize=6)
            plt.xlabel("Índice", fontsize=12)
            plt.ylabel("Valor", fontsize=12)
        elif tipo == "barras":
            plt.bar(range(len(datos_lista)), datos_lista)
            plt.xlabel("Categoría", fontsize=12)
            plt.ylabel("Valor", fontsize=12)
        else:
            plt.pie(datos_lista, autopct='%1.1f%%')
        
        plt.title(titulo, fontsize=14)
        plt.grid(True, alpha=0.3)
        plt.tight_layout()
        plt.savefig(temp_path, dpi=100, bbox_inches='tight')
        plt.close()
        
        url = f"http://localhost:8000/api/download/{filename}"
        return f"✅ Gráfica generada. **Descarga:** {url} (válido por 5 minutos)"
        
    except Exception as e:
        print(f"[TOOL] Error en generar_grafica: {e}")
        return f"Error al generar gráfica: {str(e)}"


@tool
def comparar_documentos(doc1: str, doc2: str) -> str:
    """
    Compara dos documentos y devuelve un análisis detallado.
    
    Args:
        doc1: Nombre del primer documento
        doc2: Nombre del segundo documento
    
    Returns:
        Análisis comparativo de los documentos
    
    Ejemplos:
        - comparar_documentos("contrato_2024.pdf", "contrato_2025.pdf")
        - comparar_documentos("informe_ventas.xlsx", "informe_gastos.xlsx")
    """
    print(f"[TOOL] Comparando documentos: {doc1} vs {doc2}")
    
    try:
        from backend.rag import retrieve_chunks, _get_gemini_model
        
        # Buscar chunks por nombre exacto usando metadata de ChromaDB
        try:
            results1 = collection.get(where={"filename": doc1})
            docs1 = results1.get("documents", [])[:10]
            metas1 = results1.get("metadatas", [])[:10]
        except:
            docs1, metas1 = retrieve_chunks(doc1, None, 10)
        
        try:
            results2 = collection.get(where={"filename": doc2})
            docs2 = results2.get("documents", [])[:10]
            metas2 = results2.get("metadatas", [])[:10]
        except:
            docs2, metas2 = retrieve_chunks(doc2, None, 10)
        
        if not docs1:
            return f"No se encontró el documento '{doc1}'"
        if not docs2:
            return f"No se encontró el documento '{doc2}'"
        
        texto1 = "\n".join([d[:500] for d in docs1[:5]])
        texto2 = "\n".join([d[:500] for d in docs2[:5]])
        
        model = _get_gemini_model()
        prompt = f"""
        Compara los siguientes dos documentos:
        
        DOCUMENTO 1: {doc1}
        CONTENIDO:
        {texto1[:3000]}
        
        DOCUMENTO 2: {doc2}
        CONTENIDO:
        {texto2[:3000]}
        
        Proporciona:
        1. **Similitudes** (qué tienen en común)
        2. **Diferencias** (en qué se distinguen)
        3. **Conclusión** (cuál es más útil para qué propósito)
        """
        
        response = model.generate_content(prompt)
        comparacion = response.text
        
        return f"**Comparación: {doc1} vs {doc2}**\n\n{comparacion}"
        
    except Exception as e:
        print(f"[TOOL] Error en comparar_documentos: {e}")
        return f"Error al comparar documentos: {str(e)}"


@tool
def obtener_info_sistema() -> str:
    """
    Obtiene información del sistema: documentos cargados, estadísticas.
    
    Returns:
        Información del estado del sistema
    
    Ejemplos:
        - obtener_info_sistema()
    """
    print("[TOOL] Obteniendo información del sistema")
    
    try:
        from backend.rag import get_documents_info, collection
        
        docs_info = get_documents_info()
        total_docs = docs_info["total"]
        documentos = docs_info["documentos"]
        
        total_chunks = collection.count()
        
        resultado = f"""📊 **Estado del Sistema SGD·IA**

**Documentos cargados:** {total_docs}
**Fragmentos indexados:** {total_chunks}

**Lista de documentos:**
"""
        for doc in documentos[:15]:
            resultado += f"- {doc}\n"
        
        if len(documentos) > 15:
            resultado += f"- ... y {len(documentos) - 15} más\n"
        
        resultado += f"\n**Sugerencia:** Puedes preguntar sobre cualquiera de estos documentos usando su nombre."
        
        return resultado
        
    except Exception as e:
        print(f"[TOOL] Error en obtener_info_sistema: {e}")
        return f"Error al obtener información: {str(e)}"


@tool
def obtener_historial_conversacion(user_id: int, session_id: str, limite: int = 10) -> str:
    """
    Obtiene el historial de conversación del usuario.
    
    Args:
        user_id: ID del usuario
        session_id: ID de la sesión
        limite: Número máximo de mensajes a recuperar
    
    Returns:
        Historial formateado de la conversación
    """
    print(f"[TOOL] Obteniendo historial para usuario {user_id}")
    
    try:
        contexto = get_conversation_context(user_id, session_id, limite)
        
        if not contexto:
            return "No hay historial de conversación para esta sesión."
        
        return f"**Historial de conversación:**\n\n{contexto}"
        
    except Exception as e:
        print(f"[TOOL] Error en obtener_historial_conversacion: {e}")
        return f"Error al obtener historial: {str(e)}"


# Lista de todas las herramientas para fácil acceso
HERRAMIENTAS = [
    buscar_en_documentos,
    resumir_documento,
    extraer_datos_numericos,
    generar_word,
    generar_pdf,
    generar_grafica,
    comparar_documentos,
    obtener_info_sistema,
    obtener_historial_conversacion,  # ← agregada
]